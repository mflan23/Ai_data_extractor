"""Parser service – converts uploaded files into raw text strings."""
from __future__ import annotations

import csv
import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Dispatch table
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {
    # Text
    ".txt", ".md", ".rst",
    # Documents
    ".pdf", ".docx", ".doc",
    # Spreadsheets
    ".csv", ".tsv", ".xlsx", ".xls",
    # Images
    ".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp",
}


def parse_file(filename: str, data: bytes) -> str:
    """Return extracted text for the given file bytes.

    Dispatches to the appropriate parser based on the file extension.
    Returns an empty string if the file type is unsupported.
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return _parse_pdf(data)
    if suffix in {".docx", ".doc"}:
        return _parse_docx(data)
    if suffix in {".xlsx", ".xls"}:
        return _parse_excel(data)
    if suffix in {".csv", ".tsv"}:
        delimiter = "\t" if suffix == ".tsv" else ","
        return _parse_csv(data, delimiter)
    if suffix in {".txt", ".md", ".rst"}:
        return _parse_text(data)
    if suffix in {".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp"}:
        return _parse_image(data)

    logger.warning("Unsupported file type: %s", suffix)
    return ""


# ---------------------------------------------------------------------------
# Individual parsers
# ---------------------------------------------------------------------------

def _parse_text(data: bytes) -> str:
    """Decode plain-text files with chardet fallback."""
    try:
        import chardet
        detected = chardet.detect(data)
        encoding = detected.get("encoding") or "utf-8"
    except ImportError:
        encoding = "utf-8"

    return data.decode(encoding, errors="replace").strip()


def _parse_csv(data: bytes, delimiter: str = ",") -> str:
    """Convert CSV/TSV bytes to a readable text representation."""
    try:
        import chardet
        detected = chardet.detect(data)
        encoding = detected.get("encoding") or "utf-8"
    except ImportError:
        encoding = "utf-8"

    text = data.decode(encoding, errors="replace")
    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    rows = list(reader)
    if not rows:
        return ""
    return "\n".join("\t".join(row) for row in rows)


def _parse_excel(data: bytes) -> str:
    """Convert Excel workbook bytes to text."""
    try:
        import pandas as pd
    except ImportError:
        logger.warning("pandas not installed – cannot parse Excel files")
        return ""

    parts: list[str] = []
    xl = pd.ExcelFile(io.BytesIO(data))
    for sheet_name in xl.sheet_names:
        df = xl.parse(sheet_name)
        parts.append(f"=== Sheet: {sheet_name} ===")
        parts.append(df.to_csv(index=False))
    return "\n\n".join(parts)


def _parse_pdf(data: bytes) -> str:
    """Extract text from a PDF, falling back to OCR if the layer is empty."""
    text = _pdf_text_layer(data)
    if text.strip():
        return text

    # Fall back to OCR
    logger.info("No text layer found in PDF – using OCR fallback")
    from .ocr_service import ocr_pdf_bytes
    return ocr_pdf_bytes(data)


def _pdf_text_layer(data: bytes) -> str:
    """Extract embedded text from a PDF using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber not installed – cannot parse PDF text layer")
        return ""

    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                parts.append(page_text.strip())

            # Also extract tables from the page
            for table in page.extract_tables():
                rows = ["\t".join(cell or "" for cell in row) for row in table]
                parts.append("\n".join(rows))

    return "\n\n".join(parts)


def _parse_docx(data: bytes) -> str:
    """Extract text from a Word document."""
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx not installed – cannot parse DOCX files")
        return ""

    doc = Document(io.BytesIO(data))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append("\t".join(cells))

    return "\n".join(parts)


def _parse_image(data: bytes) -> str:
    """OCR an image file."""
    from .ocr_service import ocr_image_bytes
    return ocr_image_bytes(data)
